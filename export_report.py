from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def dat_font(run, ten_font="Times New Roman", co_chu=13, dam=False):
    run.font.name = ten_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ten_font)
    run.font.size = Pt(co_chu)
    run.bold = dam


def to_mau_o(o, mau="D9EAF7"):
    tc_pr = o._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), mau)
    tc_pr.append(shd)


def can_giua_o(o):
    o.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def dinh_dang_so(gia_tri):
    try:
        so = float(gia_tri)

        if so.is_integer():
            return str(int(so))

        return f"{so:.1f}"

    except:
        return str(gia_tri)


def tao_bao_cao_word(
    nam_hoc,
    tuan,
    bang_xep_hang,
    noi_dung_nhan_xet,
    ten_truong="TRƯỜNG THCS ................................",
    dia_danh="................",
    ngay=".....",
    thang=".....",
    nam="2026",
    chuc_danh_ky="KT. HIỆU TRƯỞNG\nPHÓ HIỆU TRƯỞNG",
    nguoi_tong_hop=""
):
    tai_lieu = Document()

    # =====================================================
    # CÀI ĐẶT TRANG
    # =====================================================

    section = tai_lieu.sections[0]

    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # =====================================================
    # PHẦN ĐẦU VĂN BẢN
    # =====================================================

    bang_dau = tai_lieu.add_table(
        rows=1,
        cols=2
    )

    bang_dau.alignment = WD_TABLE_ALIGNMENT.CENTER

    o_trai = bang_dau.rows[0].cells[0]
    o_phai = bang_dau.rows[0].cells[1]

    o_trai.text = (
        f"{ten_truong}\n"
        "----------------"
    )

    o_phai.text = (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n"
        "Độc lập - Tự do - Hạnh phúc"
    )

    for p in o_trai.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in p.runs:
            dat_font(run, co_chu=12, dam=True)

    for p in o_phai.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, run in enumerate(p.runs):

            if i == 0:
                dat_font(run, co_chu=12, dam=True)
            else:
                dat_font(run, co_chu=12)

    # =====================================================
    # ĐỊA DANH - NGÀY THÁNG
    # =====================================================

    p = tai_lieu.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run(
        f"{dia_danh}, ngày {ngay} tháng {thang} năm {nam}"
    )

    dat_font(
        run,
        co_chu=12
    )

    tai_lieu.add_paragraph("")

    # =====================================================
    # TIÊU ĐỀ
    # =====================================================

    p1 = tai_lieu.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p1.add_run(
        "BẢNG THEO DÕI THI ĐUA CÁC LỚP"
    )

    dat_font(
        run,
        co_chu=16,
        dam=True
    )

    p2 = tai_lieu.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p2.add_run(
        f"NĂM HỌC {nam_hoc} - TUẦN {tuan}"
    )

    dat_font(
        run,
        co_chu=14,
        dam=True
    )

    tai_lieu.add_paragraph("")

    # =====================================================
    # PHẦN I
    # =====================================================

    p = tai_lieu.add_paragraph()

    run = p.add_run(
        "I. BẢNG XẾP HẠNG THI ĐUA"
    )

    dat_font(
        run,
        co_chu=13,
        dam=True
    )

    cac_cot = [
        "Hạng",
        "Lớp",
        "Học tập",
        "Kỷ luật",
        "Vệ sinh",
        "Điểm cộng",
        "Điểm trừ",
        "Tổng điểm",
        "Tăng/Giảm"
    ]

    bang = tai_lieu.add_table(
        rows=1,
        cols=len(cac_cot)
    )

    bang.style = "Table Grid"
    bang.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Tiêu đề bảng
    for i, ten_cot in enumerate(cac_cot):

        o = bang.rows[0].cells[i]

        o.text = ten_cot

        to_mau_o(
            o,
            "D9EAF7"
        )

        can_giua_o(o)

        for p in o.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in p.runs:
                dat_font(
                    run,
                    co_chu=10,
                    dam=True
                )

    # Dữ liệu bảng
    for _, dong in bang_xep_hang.iterrows():

        hang_moi = bang.add_row().cells

        du_lieu_dong = [
            dong.get("Hạng", ""),
            dong.get("Lớp", ""),
            dinh_dang_so(dong.get("Học tập", "")),
            dinh_dang_so(dong.get("Kỷ luật", "")),
            dinh_dang_so(dong.get("Vệ sinh", "")),
            dinh_dang_so(dong.get("Điểm cộng", "")),
            dinh_dang_so(dong.get("Điểm trừ", "")),
            dinh_dang_so(dong.get("Tổng điểm", "")),
            dong.get("Tăng/Giảm", "")
        ]

        for i, gia_tri in enumerate(du_lieu_dong):

            hang_moi[i].text = str(gia_tri)

            can_giua_o(
                hang_moi[i]
            )

            for p in hang_moi[i].paragraphs:

                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in p.runs:
                    dat_font(
                        run,
                        co_chu=10
                    )

    tai_lieu.add_paragraph("")

    # =====================================================
    # PHẦN II
    # =====================================================

    p = tai_lieu.add_paragraph()

    run = p.add_run(
        "II. NHẬN XÉT THI ĐUA TRONG TUẦN"
    )

    dat_font(
        run,
        co_chu=13,
        dam=True
    )

    if noi_dung_nhan_xet:

        cac_doan = noi_dung_nhan_xet.split("\n\n")

        for noi_dung in cac_doan:

            if noi_dung.strip():

                p = tai_lieu.add_paragraph()

                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                p.paragraph_format.first_line_indent = Cm(1)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)

                run = p.add_run(
                    noi_dung.strip()
                )

                dat_font(
                    run,
                    co_chu=13
                )

    else:

        p = tai_lieu.add_paragraph()

        run = p.add_run(
            "Chưa có nội dung nhận xét."
        )

        dat_font(
            run,
            co_chu=13
        )

    tai_lieu.add_paragraph("")

    # =====================================================
    # PHẦN KÝ TÊN
    # =====================================================

    bang_ky = tai_lieu.add_table(
        rows=1,
        cols=2
    )

    bang_ky.alignment = WD_TABLE_ALIGNMENT.CENTER

    o1 = bang_ky.rows[0].cells[0]
    o2 = bang_ky.rows[0].cells[1]

    o1.text = (
        "NGƯỜI TỔNG HỢP\n"
        "(Ký, ghi rõ họ tên)\n\n\n"
        f"{nguoi_tong_hop}"
    )

    o2.text = (
        f"{chuc_danh_ky}\n"
        "(Ký, ghi rõ họ tên)\n\n\n\n"
    )

    for o in [o1, o2]:

        can_giua_o(o)

        for p in o.paragraphs:

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for i, run in enumerate(p.runs):

                if i == 0:
                    dat_font(
                        run,
                        co_chu=12,
                        dam=True
                    )
                else:
                    dat_font(
                        run,
                        co_chu=11
                    )

    # =====================================================
    # LƯU FILE
    # =====================================================

    ten_file = (
        f"reports/bao_cao_thi_dua_"
        f"tuan_{tuan}.docx"
    )

    tai_lieu.save(
        ten_file
    )

    return ten_file