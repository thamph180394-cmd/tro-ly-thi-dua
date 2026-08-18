import os
import pandas as pd

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# HÀM HỖ TRỢ
# ============================================================

def hien_thi_so(gia_tri):
    try:
        so = float(gia_tri)

        if so.is_integer():
            return str(int(so))

        return f"{so:.2f}".rstrip("0").rstrip(".")

    except Exception:
        return str(gia_tri)


def dat_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Times New Roman"
    )


def dat_le_trang(section):
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def to_mau_o(cell, mau="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), mau)

    tc_pr.append(shd)


def can_giua_o(cell):
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    for paragraph in cell.paragraphs:
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )


def dat_duong_vien_o(cell, mau="A6A6A6"):
    tc_pr = cell._tc.get_or_add_tcPr()

    borders = (
        tc_pr.first_child_found_in(
            "w:tcBorders"
        )
    )

    if borders is None:
        borders = OxmlElement(
            "w:tcBorders"
        )
        tc_pr.append(borders)

    for edge in [
        "top",
        "left",
        "bottom",
        "right"
    ]:

        tag = "w:" + edge

        element = borders.find(
            qn(tag)
        )

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(
            qn("w:val"),
            "single"
        )

        element.set(
            qn("w:sz"),
            "4"
        )

        element.set(
            qn("w:space"),
            "0"
        )

        element.set(
            qn("w:color"),
            mau
        )


def xoa_vien_bang(table):
    for row in table.rows:
        for cell in row.cells:

            tc_pr = (
                cell._tc.get_or_add_tcPr()
            )

            borders = OxmlElement(
                "w:tcBorders"
            )

            for edge in [
                "top",
                "left",
                "bottom",
                "right",
                "insideH",
                "insideV"
            ]:

                element = OxmlElement(
                    "w:" + edge
                )

                element.set(
                    qn("w:val"),
                    "nil"
                )

                borders.append(element)

            tc_pr.append(borders)


def ghep_danh_sach_lop(ds_lop):
    ds_lop = [
        str(x)
        for x in ds_lop
    ]

    if len(ds_lop) == 0:
        return ""

    if len(ds_lop) == 1:
        return ds_lop[0]

    if len(ds_lop) == 2:
        return (
            ds_lop[0]
            + " và "
            + ds_lop[1]
        )

    return (
        ", ".join(ds_lop[:-1])
        + " và "
        + ds_lop[-1]
    )


# ============================================================
# CHUẨN HÓA + XẾP HẠNG
# ============================================================

def chuan_hoa_bang(bang):
    if bang is None:
        return pd.DataFrame()

    df = bang.copy()

    if df.empty:
        return df

    if "Lớp" not in df.columns:
        raise ValueError(
            "Không tìm thấy cột Lớp."
        )

    if "Tổng điểm" not in df.columns:
        raise ValueError(
            "Không tìm thấy cột Tổng điểm."
        )

    cac_cot_so = [
        "Học tập",
        "Kỷ luật",
        "Vệ sinh",
        "Điểm cộng",
        "Điểm trừ",
        "Tổng điểm"
    ]

    for cot in cac_cot_so:
        if cot in df.columns:
            df[cot] = pd.to_numeric(
                df[cot],
                errors="coerce"
            ).fillna(0)

    df = df.sort_values(
        by=[
            "Tổng điểm",
            "Lớp"
        ],
        ascending=[
            False,
            True
        ],
        kind="stable"
    ).reset_index(drop=True)

    # Đồng điểm = đồng hạng
    # Ví dụ 1, 1, 3
    df["Hạng"] = (
        df["Tổng điểm"]
        .rank(
            method="min",
            ascending=False
        )
        .astype(int)
    )

    return df


# ============================================================
# TIÊU ĐỀ
# ============================================================

def them_tieu_de(
    doc,
    nam_hoc,
    tuan,
    ten_truong,
    dia_danh,
    ngay,
    thang,
    nam
):

    bang_dau = doc.add_table(
        rows=1,
        cols=2
    )

    bang_dau.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    # BÊN TRÁI
    cell = bang_dau.cell(
        0,
        0
    )

    p = cell.paragraphs[0]
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        str(ten_truong).upper()
    )

    dat_font(
        r,
        11,
        True
    )

    p = cell.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "LIÊN ĐỘI TH&THCS AN LINH"
    )

    dat_font(
        r,
        11,
        True
    )

    # BÊN PHẢI
    cell = bang_dau.cell(
        0,
        1
    )

    p = cell.paragraphs[0]
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    )

    dat_font(
        r,
        11,
        True
    )

    p = cell.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "Độc lập - Tự do - Hạnh phúc"
    )

    dat_font(
        r,
        11,
        True
    )

    p = cell.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        f"{dia_danh}, ngày {ngay} "
        f"tháng {thang} năm {nam}"
    )

    dat_font(
        r,
        11,
        False
    )

    xoa_vien_bang(
        bang_dau
    )

    # TIÊU ĐỀ CHÍNH
    p = doc.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

    r = p.add_run(
        "BẢNG TỔNG KẾT THI ĐUA CÁC LỚP"
    )

    dat_font(
        r,
        16,
        True
    )

    p = doc.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)

    r = p.add_run(
        f"Năm học {nam_hoc} - Tuần {tuan}"
    )

    dat_font(
        r,
        13,
        True
    )


# ============================================================
# BẢNG TOÀN TRƯỜNG
# ============================================================

def them_bang_xep_hang(
    doc,
    df
):

    p = doc.add_paragraph()

    r = p.add_run(
        "I. BẢNG XẾP HẠNG THI ĐUA TOÀN TRƯỜNG"
    )

    dat_font(
        r,
        12,
        True
    )

    cac_cot = [
        "Hạng",
        "Lớp",
        "Khối",
        "Học tập",
        "Kỷ luật",
        "Vệ sinh",
        "Điểm cộng",
        "Điểm trừ",
        "Tổng điểm"
    ]

    cac_cot = [
        cot
        for cot in cac_cot
        if cot in df.columns
    ]

    table = doc.add_table(
        rows=1,
        cols=len(cac_cot)
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = True

    # HEADER
    for i, cot in enumerate(cac_cot):

        cell = table.rows[0].cells[i]
        cell.text = cot

        to_mau_o(
            cell,
            "D9EAF7"
        )

        can_giua_o(
            cell
        )

        dat_duong_vien_o(
            cell
        )

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                dat_font(
                    run,
                    9,
                    True
                )

    # DỮ LIỆU
    for _, row in df.iterrows():

        cells = (
            table.add_row().cells
        )

        for i, cot in enumerate(cac_cot):

            value = row[cot]

            if cot not in [
                "Lớp",
                "Khối"
            ]:
                value = hien_thi_so(
                    value
                )

            cells[i].text = str(value)

            can_giua_o(
                cells[i]
            )

            dat_duong_vien_o(
                cells[i]
            )

            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    dat_font(
                        run,
                        9,
                        False
                    )


# ============================================================
# TOP 5
# ============================================================

def them_top5(
    doc,
    df
):

    p = doc.add_paragraph()

    p.paragraph_format.space_before = Pt(8)

    r = p.add_run(
        "II. TOP 5 THI ĐUA TOÀN TRƯỜNG"
    )

    dat_font(
        r,
        12,
        True
    )

    top5 = df[
        df["Hạng"] <= 5
    ].copy()

    cac_cot = [
        "Hạng",
        "Lớp",
        "Khối",
        "Tổng điểm",
        "Ghi chú"
    ]

    table = doc.add_table(
        rows=1,
        cols=len(cac_cot)
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = True

    # HEADER
    for i, cot in enumerate(cac_cot):

        cell = table.rows[0].cells[i]
        cell.text = cot

        to_mau_o(
            cell,
            "F4CCCC"
        )

        can_giua_o(
            cell
        )

        dat_duong_vien_o(
            cell
        )

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                dat_font(
                    run,
                    9,
                    True
                )

    # DỮ LIỆU
    for _, row in top5.iterrows():

        hang = int(
            row["Hạng"]
        )

        if hang == 1:
            ghi_chu = "Dẫn đầu"

        elif hang <= 3:
            ghi_chu = "Top 3"

        else:
            ghi_chu = ""

        khoi = ""

        if "Khối" in row.index:
            khoi = row["Khối"]

        values = [
            hien_thi_so(
                row["Hạng"]
            ),
            str(
                row["Lớp"]
            ),
            str(
                khoi
            ),
            hien_thi_so(
                row["Tổng điểm"]
            ),
            ghi_chu
        ]

        cells = (
            table.add_row().cells
        )

        for i, value in enumerate(values):

            cells[i].text = str(value)

            can_giua_o(
                cells[i]
            )

            dat_duong_vien_o(
                cells[i]
            )

            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    dat_font(
                        run,
                        9,
                        False
                    )


# ============================================================
# TẠO NHẬN XÉT CHUẨN TỪ BẢNG
# ============================================================

def tao_nhan_xet_chuan(
    df,
    tuan
):

    cac_doan = []

    tong_lop = len(df)

    diem_tb = (
        df["Tổng điểm"].mean()
    )

    cac_doan.append(
        f"Qua tổng hợp kết quả thi đua tuần {tuan}, "
        f"có {tong_lop} lớp được theo dõi và xếp hạng. "
        f"Điểm trung bình của các lớp đạt "
        f"{diem_tb:.2f} điểm."
    )

    # ========================================================
    # DẪN ĐẦU - XỬ LÝ ĐỒNG HẠNG
    # ========================================================

    hang_dau = df[
        df["Hạng"] == 1
    ].copy()

    ds_dau = (
        hang_dau["Lớp"]
        .astype(str)
        .tolist()
    )

    ten_dau = ghep_danh_sach_lop(
        ds_dau
    )

    diem_dau = (
        hang_dau.iloc[0]["Tổng điểm"]
    )

    if len(ds_dau) == 1:

        cac_doan.append(
            f"Lớp {ten_dau} dẫn đầu bảng xếp hạng "
            f"với {hien_thi_so(diem_dau)} điểm. "
            f"Đây là tập thể có kết quả thi đua nổi bật "
            f"trong tuần."
        )

    else:

        cac_doan.append(
            f"Các lớp {ten_dau} cùng dẫn đầu bảng xếp hạng "
            f"với {hien_thi_so(diem_dau)} điểm. "
            f"Đây là các tập thể có kết quả thi đua nổi bật "
            f"trong tuần."
        )

    # ========================================================
    # HỌC TẬP
    # ========================================================

    if "Học tập" in df.columns:

        max_diem = (
            df["Học tập"].max()
        )

        ds = (
            df[
                df["Học tập"]
                == max_diem
            ]["Lớp"]
            .astype(str)
            .tolist()
        )

        ten = ghep_danh_sach_lop(
            ds
        )

        if len(ds) == 1:

            cac_doan.append(
                f"Về học tập, lớp {ten} có điểm học tập "
                f"cao nhất với {hien_thi_so(max_diem)} điểm."
            )

        else:

            cac_doan.append(
                f"Về học tập, các lớp {ten} cùng có điểm "
                f"học tập cao nhất với "
                f"{hien_thi_so(max_diem)} điểm."
            )

    # ========================================================
    # KỶ LUẬT
    # ========================================================

    if "Kỷ luật" in df.columns:

        max_diem = (
            df["Kỷ luật"].max()
        )

        ds = (
            df[
                df["Kỷ luật"]
                == max_diem
            ]["Lớp"]
            .astype(str)
            .tolist()
        )

        ten = ghep_danh_sach_lop(
            ds
        )

        if len(ds) == 1:

            cac_doan.append(
                f"Về kỷ luật, lớp {ten} có điểm kỷ luật "
                f"cao nhất với {hien_thi_so(max_diem)} điểm."
            )

        else:

            cac_doan.append(
                f"Về kỷ luật, các lớp {ten} cùng có điểm "
                f"kỷ luật cao nhất với "
                f"{hien_thi_so(max_diem)} điểm."
            )

    # ========================================================
    # VỆ SINH
    # ========================================================

    if "Vệ sinh" in df.columns:

        max_diem = (
            df["Vệ sinh"].max()
        )

        ds = (
            df[
                df["Vệ sinh"]
                == max_diem
            ]["Lớp"]
            .astype(str)
            .tolist()
        )

        ten = ghep_danh_sach_lop(
            ds
        )

        if len(ds) == 1:

            cac_doan.append(
                f"Về vệ sinh, lớp {ten} có điểm vệ sinh "
                f"cao nhất với {hien_thi_so(max_diem)} điểm."
            )

        else:

            cac_doan.append(
                f"Về vệ sinh, các lớp {ten} cùng có điểm "
                f"vệ sinh cao nhất với "
                f"{hien_thi_so(max_diem)} điểm."
            )

    # ========================================================
    # TIẾN BỘ / GIẢM HẠNG
    # ========================================================

    if "Tăng/Giảm" in df.columns:

        tang_max = 0
        lop_tang = None

        giam_max = 0
        lop_giam = None

        for _, row in df.iterrows():

            text = str(
                row["Tăng/Giảm"]
            )

            if "▲" in text:

                try:

                    so = int(
                        text.replace(
                            "▲",
                            ""
                        ).strip()
                    )

                    if so > tang_max:

                        tang_max = so
                        lop_tang = row["Lớp"]

                except Exception:
                    pass

            if "▼" in text:

                try:

                    so = int(
                        text.replace(
                            "▼",
                            ""
                        ).strip()
                    )

                    if so > giam_max:

                        giam_max = so
                        lop_giam = row["Lớp"]

                except Exception:
                    pass

        if lop_tang is not None:

            cac_doan.append(
                f"Lớp {lop_tang} là tập thể có sự tiến bộ "
                f"nổi bật nhất khi tăng {tang_max} bậc "
                f"so với tuần trước."
            )

        if lop_giam is not None:

            cac_doan.append(
                f"Lớp {lop_giam} giảm {giam_max} bậc "
                f"so với tuần trước. Tập thể lớp cần rà soát "
                f"các nội dung còn hạn chế để có biện pháp "
                f"khắc phục trong tuần tiếp theo."
            )

    # ========================================================
    # ĐIỂM TRỪ CAO NHẤT
    # ========================================================

    if "Điểm trừ" in df.columns:

        max_tru = (
            df["Điểm trừ"].max()
        )

        ds = (
            df[
                df["Điểm trừ"]
                == max_tru
            ]["Lớp"]
            .astype(str)
            .tolist()
        )

        ten = ghep_danh_sach_lop(
            ds
        )

        if len(ds) == 1:

            cac_doan.append(
                f"Lớp {ten} có số điểm trừ cao nhất "
                f"trong tuần với {hien_thi_so(max_tru)} điểm. "
                f"Đề nghị tập thể lớp chú ý hạn chế các vi phạm "
                f"và nâng cao ý thức thực hiện nội quy."
            )

        else:

            cac_doan.append(
                f"Các lớp {ten} cùng có số điểm trừ cao nhất "
                f"trong tuần với {hien_thi_so(max_tru)} điểm. "
                f"Đề nghị các tập thể chú ý hạn chế các vi phạm "
                f"và nâng cao ý thức thực hiện nội quy."
            )

    # ========================================================
    # CUỐI BẢNG
    # ========================================================

    diem_cuoi = (
        df["Tổng điểm"].min()
    )

    ds_cuoi = (
        df[
            df["Tổng điểm"]
            == diem_cuoi
        ]["Lớp"]
        .astype(str)
        .tolist()
    )

    ten_cuoi = ghep_danh_sach_lop(
        ds_cuoi
    )

    if len(ds_cuoi) == 1:

        cac_doan.append(
            f"Lớp {ten_cuoi} hiện xếp cuối bảng với "
            f"{hien_thi_so(diem_cuoi)} điểm. "
            f"Tập thể lớp cần tiếp tục cố gắng, đặc biệt "
            f"chú trọng nâng cao kết quả học tập, ý thức "
            f"kỷ luật, vệ sinh và hạn chế điểm trừ."
        )

    else:

        cac_doan.append(
            f"Các lớp {ten_cuoi} hiện cùng xếp cuối bảng với "
            f"{hien_thi_so(diem_cuoi)} điểm. "
            f"Các tập thể cần tiếp tục cố gắng, đặc biệt "
            f"chú trọng nâng cao kết quả học tập, ý thức "
            f"kỷ luật, vệ sinh và hạn chế điểm trừ."
        )

    # ========================================================
    # KẾT
    # ========================================================

    cac_doan.append(
        "Đề nghị các lớp tiếp tục phát huy những mặt đã "
        "thực hiện tốt, duy trì tinh thần thi đua tích cực, "
        "chấp hành nghiêm nội quy, nâng cao chất lượng học tập "
        "và giữ gìn vệ sinh lớp học. Các tập thể còn hạn chế "
        "cần chủ động khắc phục để nâng cao kết quả trong "
        "tuần tiếp theo."
    )

    return cac_doan


# ============================================================
# NHẬN XÉT
# ============================================================

def them_nhan_xet(
    doc,
    df,
    tuan,
    noi_dung_nhan_xet
):

    p = doc.add_paragraph()

    p.paragraph_format.space_before = Pt(
        8
    )

    r = p.add_run(
        "III. NHẬN XÉT THI ĐUA TRONG TUẦN"
    )

    dat_font(
        r,
        12,
        True
    )

    # ========================================================
    # QUAN TRỌNG:
    # KHÔNG dùng nhận xét AI cũ nữa.
    #
    # Tạo lại nhận xét trực tiếp từ bảng đã xếp hạng
    # để đồng hạng luôn chính xác.
    # ========================================================

    cac_doan = tao_nhan_xet_chuan(
        df,
        tuan
    )

    for noi_dung in cac_doan:

        p = doc.add_paragraph()

        p.alignment = (
            WD_ALIGN_PARAGRAPH.JUSTIFY
        )

        p.paragraph_format.first_line_indent = Cm(
            1
        )

        p.paragraph_format.space_after = Pt(
            3
        )

        r = p.add_run(
            noi_dung
        )

        dat_font(
            r,
            12,
            False
        )


# ============================================================
# PHẦN KÝ
# ============================================================

def them_phan_ky(
    doc,
    chuc_danh_ky,
    nguoi_tong_hop
):

    doc.add_paragraph()

    table = doc.add_table(
        rows=1,
        cols=2
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    # NGƯỜI TỔNG HỢP
    cell = table.cell(
        0,
        0
    )

    p = cell.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "NGƯỜI TỔNG HỢP"
    )

    dat_font(
        r,
        12,
        True
    )

    p = cell.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "(Ký, ghi rõ họ tên)"
    )

    dat_font(
        r,
        11,
        False
    )

    for _ in range(3):
        cell.add_paragraph()

    if str(
        nguoi_tong_hop
    ).strip():

        p = cell.add_paragraph()

        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        r = p.add_run(
            str(
                nguoi_tong_hop
            )
        )

        dat_font(
            r,
            12,
            True
        )

    # NGƯỜI KÝ
    cell = table.cell(
        0,
        1
    )

    cac_dong = str(
        chuc_danh_ky
    ).split("\n")

    for i, dong in enumerate(cac_dong):

        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()

        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        r = p.add_run(
            dong
        )

        dat_font(
            r,
            12,
            True
        )

    p = cell.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "(Ký, ghi rõ họ tên)"
    )

    dat_font(
        r,
        11,
        False
    )

    xoa_vien_bang(
        table
    )


# ============================================================
# HÀM CHÍNH
#
# NHẬN ĐÚNG 11 THAM SỐ TỪ APP.PY
# ============================================================

def tao_bao_cao_word_toan_truong(
    nam_hoc,
    tuan,
    bang,
    noi_dung_nhan_xet="",
    ten_truong="TRƯỜNG TH&THCS AN LINH",
    dia_danh="An Linh",
    ngay=".....",
    thang=".....",
    nam="2026",
    chuc_danh_ky="KT. HIỆU TRƯỞNG\nPHÓ HIỆU TRƯỞNG",
    nguoi_tong_hop=""
):

    df = chuan_hoa_bang(
        bang
    )

    if df.empty:
        return None

    doc = Document()

    for section in doc.sections:

        section.page_width = Cm(
            21
        )

        section.page_height = Cm(
            29.7
        )

        dat_le_trang(
            section
        )

    # Font mặc định
    style = doc.styles[
        "Normal"
    ]

    style.font.name = (
        "Times New Roman"
    )

    style.font.size = Pt(
        12
    )

    style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Times New Roman"
    )

    # Nội dung
    them_tieu_de(
        doc,
        nam_hoc,
        tuan,
        ten_truong,
        dia_danh,
        ngay,
        thang,
        nam
    )

    them_bang_xep_hang(
        doc,
        df
    )

    them_top5(
        doc,
        df
    )

    them_nhan_xet(
        doc,
        df,
        tuan,
        noi_dung_nhan_xet
    )

    them_phan_ky(
        doc,
        chuc_danh_ky,
        nguoi_tong_hop
    )

    # Lưu
    ten_file = (
        f"bao_cao_toan_truong_tuan_{tuan}.docx"
    )

    duong_dan = os.path.abspath(
        ten_file
    )

    doc.save(
        duong_dan
    )

    return duong_dan


# ============================================================
# HÀM DỰ PHÒNG
# ============================================================

def tao_word_toan_truong(
    nam_hoc,
    tuan,
    bang
):

    return tao_bao_cao_word_toan_truong(
        nam_hoc,
        tuan,
        bang
    )


def xuat_word_toan_truong(
    nam_hoc,
    tuan,
    bang
):

    return tao_bao_cao_word_toan_truong(
        nam_hoc,
        tuan,
        bang
    )